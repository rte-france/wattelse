#  Copyright (c) 2024, RTE (https://www.rte-france.com)
#  See AUTHORS.txt
#  SPDX-License-Identifier: MPL-2.0
#  This file is part of Wattelse, a NLP application suite.

import typer
import pandas as pd
import numpy as np
from loguru import logger
from typing import List, Dict
from tqdm import tqdm
from docx import Document
from collections import defaultdict
from pathlib import Path
import PyPDF2
from wattelse.api.openai.client_openai_api import OpenAI_Client 
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from concurrent.futures import ThreadPoolExecutor
import re

# TODO REFACTOR : Use the classes rather than creating the functions.
# from wattelse.indexer.document_parser import parse_file
# from wattelse.indexer.document_splitter import split_file

from wattelse.chatbot.eval.prompt import (
    QA_GENERATION_PROMPT,
    QUESTION_GROUNDEDNESS_CRITIQUE_PROMPT,
    QUESTION_REALISM_CRITIQUE_PROMPT,
    QUESTION_STANDALONE_CRITIQUE_PROMPT
)

# Example :
#  python generate_qa.py --eval-corpus-path data/eval_CCRT --n-generations 10 --output-path output_qa_201-CCRT.xlsx --report-output-path report_output_201-CCRT.xlsx

# Define the Typer app
app = typer.Typer()

# Function to call the LLM (Refactor with classes WattElse)
def call_llm(llm_client, prompt: str) -> str:
    response = llm_client.generate(prompt, temperature=0)
    return response

# Enhanced split_documents function for parallel processing
def split_documents(eval_corpus_path: Path) -> List[LangchainDocument]:
    docs_processed = []

    def process_doc(doc):
        content = ""
        if doc.suffix == ".pdf":
            try:
                with open(doc, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text() is not None])
            except Exception as e:
                logger.error(f"Error reading {doc.name}: {e}")
                return []
        elif doc.suffix == ".docx":
            try:
                docx_file = Document(doc)
                content = "\n".join([paragraph.text for paragraph in docx_file.paragraphs if paragraph.text.strip() != ""])
            except Exception as e:
                logger.error(f"Error reading {doc.name}: {e}")
                return []
        else:
            try:
                with open(doc, "r", encoding="utf-8") as f:
                    content = f.read()
            except Exception as e:
                logger.error(f"Error reading {doc.name}: {e}")
                return []

        if content:
            logger.info(f"Processing document '{doc.name}' with total length: {len(content)} characters")
            langchain_doc = LangchainDocument(page_content=content, metadata={"source": doc.name})
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=3000,
                chunk_overlap=200,
                add_start_index=True, # This enables assigning start indexes to chunks
                separators=["\n\n", "\n", ".", " ", ""],
            )
            return text_splitter.split_documents([langchain_doc])
        return []

    # To handle multiple documents efficiently, the splitting process is parallelized: ThreadPoolExecutor()
    with ThreadPoolExecutor() as executor:
        all_chunks = list(executor.map(process_doc, eval_corpus_path.iterdir()))
        for chunks in all_chunks:
            docs_processed.extend(chunks) # Chunks are gathered into a single list

    return docs_processed

# Main QA generation function, default parameters.
def generate_qa_pairs(
    EVAL_CORPUS_PATH: Path,
    N_GENERATIONS: int = 100,
    OUTPUT_PATH: Path = Path("qa_output.xlsx"),
    DOCS_PER_QA: int = 1,
    CHUNKS_PER_DOC: int = 1
) -> List[Dict]:
    '''
    Function to generate QA pairs from documents. It processes documents, selects chunks,
    and generates QA pairs using an LLM. 

    Parameters:
    - EVAL_CORPUS_PATH: Path to the folder containing the documents to be processed.
    - N_GENERATIONS: The number of QA pairs to generate (default is 100).
    - OUTPUT_PATH: The path to save the generated QA pairs as an Excel file.
    - DOCS_PER_QA: The number of documents to use for each QA generation.
    - CHUNKS_PER_DOC: The number of chunks to use from each document for a single QA generation.

    Returns:
    - A list of dictionaries containing generated QA pairs.
    '''

    outputs = []
    # Step 1: Split the documents into chunks
    docs_processed = split_documents(EVAL_CORPUS_PATH)

    # Step 2: Organize chunks by document
    chunks_by_doc = defaultdict(list)
    for chunk in docs_processed:
        source = chunk.metadata["source"]
        chunks_by_doc[source].append(chunk)

    # Step 3: Track chunk usage for each document
    doc_names = list(chunks_by_doc.keys()) # List of document names (not used here)
    chunk_usage = {doc_name: [False] * len(chunks) for doc_name, chunks in chunks_by_doc.items()}

    # Step 4: Initialize the LLM client for generating QA pairs
    llm_client = OpenAI_Client() 
    logger.info(f"LLM Generation model: {llm_client.model_name}")

    # Step 5 : Function to select a high-yield document based on chunk availability. It's basically a ratio.
    def select_high_yield_document() -> List[str]:
        '''
        Selects documents that have the most available chunks for QA generation.
        
        Returns:
        - A sorted list of document names based on the number of available chunks.
        '''
        doc_chunks_remaining = {
            doc: sum(1 for used in usage if not used) / len(usage)
            for doc, usage in chunk_usage.items()
        }
        # Sort documents by the proportion of available chunks (descending order)
        high_yield_docs = sorted(doc_chunks_remaining, key=doc_chunks_remaining.get, reverse=True)
        return high_yield_docs

    # Step 6: Calculate the maximum number of possible generations based on available chunks.
    # TODO: Review the generation method. 
    max_generations = 0
    for doc, chunks in chunks_by_doc.items():
        available_chunks_count = sum(not used for used in chunk_usage[doc])
        max_generations += available_chunks_count // CHUNKS_PER_DOC

    # Ensure the number of generations does not exceed the available chunks
    actual_generations = min(N_GENERATIONS, max_generations)

    # Step 7: Loop for QA generation
    for iteration in tqdm(range(actual_generations), total=actual_generations):
        '''
        Loop through the required number of generations. For each iteration:
        - Select high-yield documents.
        - Choose chunks from the selected documents.
        - Generate QA pairs using the LLM.
        '''

        sampled_contexts = [] # List to store chunks for the current QA pair generation
        used_docs = set() # Set to track used documents in the current generation
        logger.info(f"Iteration {iteration + 1}/{actual_generations}")

        # Try selecting high-yield documents first
        # The code calls select_high_yield_document() to get the list of documents with the highest available chunks. 
        # It then selects documents (one per iteration) that haven't been used yet in the current generation.

        for _ in range(DOCS_PER_QA):
            high_yield_docs = select_high_yield_document() # Select documents with the most available chunks
            selected_doc = None # Variable to store the selected document

            # Step 8: Loop through documents to select one that hasn't been used yet in the current iteration
            for doc_name in high_yield_docs:
                if doc_name not in used_docs:
                    selected_doc = doc_name
                    used_docs.add(doc_name)
                    break

            # Step 9: Handle case when no new documents are available
            if selected_doc is None:
                logger.info("No new documents available. Ending QA generation.")
                break  # Break out if no new documents are available, making sure to avoid repetitions of chunks.

            # Step 10: Get chunks from the selected document
            doc_chunks = chunks_by_doc[selected_doc]
            available_chunks = [idx for idx, used in enumerate(chunk_usage[selected_doc]) if not used]

            # Step 11: Select chunks from the available ones
            if available_chunks:
                chunk_found = 0
                for chunk_idx in available_chunks:
                    if chunk_found < CHUNKS_PER_DOC:
                        chunk_usage[selected_doc][chunk_idx] = True  # Mark chunk as used
                        sampled_contexts.append(doc_chunks[chunk_idx]) # Add to context
                        logger.info(f"Using chunk from {selected_doc}: [{chunk_idx}]") # Chunks used from the selected docs
                        chunk_found += 1
            else:
                logger.info(f"All chunks exhausted for document '{selected_doc}'. Skipping this round.")

        # Step 12: Handle case when no contexts are available for QA generation
        if not sampled_contexts:
            logger.warning("No available contexts. Ending QA generation process.")
            break  # Break out if no contexts are available, avoiding empty generations.

        # Step 13: Combine the selected contexts and call the LLM to generate QA pairs
        combined_context = "\n\n".join(chunk.page_content for chunk in sampled_contexts)
        output_QA_couple = call_llm(llm_client, QA_GENERATION_PROMPT.format(context=combined_context))

        try:
            # Step 14: Extract the generated QA pairs from the response
            simple_question_match = re.search(r"\*?\*?Question simple\*?\*?\s*:\s*(.*)", output_QA_couple)
            simple_answer_match = re.search(r"\*?\*?Réponse simple\*?\*?\s*:\s*(.*)", output_QA_couple)
            reasoning_question_match = re.search(r"\*?\*?Question de raisonnement\*?\*?\s*:\s*(.*)", output_QA_couple)
            reasoning_answer_match = re.search(r"\*?\*?Réponse de raisonnement\*?\*?\s*:\s*(.*)", output_QA_couple)
            multi_context_question_match = re.search(r"\*?\*?Question multi-contexte\*?\*?\s*:\s*(.*)", output_QA_couple)
            multi_context_answer_match = re.search(r"\*?\*?Réponse multi-contexte\*?\*?\s*:\s*(.*)", output_QA_couple)

            # Use .group(1) if the match is found; otherwise, default to an empty string or a suitable placeholder
            simple_question = simple_question_match.group(1).split("\n")[0].strip() if simple_question_match else ""
            simple_answer = simple_answer_match.group(1).split("\n")[0].strip() if simple_answer_match else ""
            reasoning_question = reasoning_question_match.group(1).split("\n")[0].strip() if reasoning_question_match else ""
            reasoning_answer = reasoning_answer_match.group(1).split("\n")[0].strip() if reasoning_answer_match else ""
            multi_context_question = multi_context_question_match.group(1).split("\n")[0].strip() if multi_context_question_match else ""
            multi_context_answer = multi_context_answer_match.group(1).split("\n")[0].strip() if multi_context_answer_match else ""


            # Step 15: Ensure answer length constraints are respected
            assert len(simple_answer) < 2000, "Answer is too long"
            assert len(reasoning_answer) < 2000, "Answer is too long"
            assert len(multi_context_answer) < 2000, "Answer is too long"

            # Step 16: Append the generated QA pair to the output list
            source_docs = [chunk.metadata["source"] for chunk in sampled_contexts]
            outputs.append(
                {
                    "context": combined_context,
                    "questions": {
                        "simple": simple_question,
                        "reasoning": reasoning_question,
                        "multi_context": multi_context_question,
                    },
                    "answers": {
                        "simple": simple_answer,
                        "reasoning": reasoning_answer,
                        "multi_context": multi_context_answer,
                    },
                    "source_docs": source_docs,
                }
            )
        except Exception as e:
            logger.error(f"Error generating QA pair: {e}")  # Log any errors encountered
            continue

    # Step 17: Save the generated QA pairs to an Excel file
    qa_df = pd.DataFrame(outputs)
    qa_df.to_excel(OUTPUT_PATH, index=False)
    logger.info(f"Generated QA dataset saved to {OUTPUT_PATH}")

    # Document utilization summary for debugging and further adjustments
    logger.info("Document Utilization Summary:")
    for doc_name, usage in chunk_usage.items():
        total_chunks = len(usage)
        used_chunks = sum(usage)
        logger.info(f"{doc_name}: {used_chunks}/{total_chunks} chunks used ({(used_chunks/total_chunks)*100:.2f}% utilization)")

    return outputs

# Function to evaluate QA pairs using critique agents
def evaluate_qa_pairs(outputs: List[Dict]) -> List[Dict]:
    '''
    Function to evaluate QA pairs generated from the LLM using critique agents
    It evaluates three types of questions: simple, reasoning, and multi-context
    and calculates groundedness, realism, and standalone quality scores.
    '''
    llm_client = OpenAI_Client()  # Initialize the OpenAI Client for critique
    logger.info(f"LLM Evaluation model: {llm_client.model_name}")

    logger.info("Generating critique for each QA couple...")
    for output in tqdm(outputs):
        try:
            # Initialize evaluation data for this QA pair
            evaluations = {}
            for complexity in ["simple", "reasoning", "multi_context"]:
                question = output["questions"][complexity]
                context = output["context"]

                # Evaluate groundedness
                groundedness_eval = call_llm(
                    llm_client,
                    QUESTION_GROUNDEDNESS_CRITIQUE_PROMPT.format(context=context, question=question),
                )
                
                complexity_groundedness_eval_match = re.search(
                    rf"^(.*?Note totale\s*:\s*([1-5]))", 
                    groundedness_eval, 
                    re.DOTALL
                )

                # Check if the match was successful
                if complexity_groundedness_eval_match:
                    evaluation_text = complexity_groundedness_eval_match.group(1).strip()  # Everything up to and including "Note totale : score"
                    score_text = complexity_groundedness_eval_match.group(2).strip()       # Just the score

                    evaluations[f"{complexity}_groundedness"] = evaluation_text
                    evaluations[f"{complexity}_groundedness_score"] = int(score_text)
                else:
                    evaluations[f"{complexity}_groundedness_score"] = np.nan

                logger.debug(f"groundedness LLM response: {groundedness_eval}")

                # # Extract evaluations and scores for realism
                # evaluations[f"{complexity}_realism"] = realism_eval.split("Évaluation : ")[-1].split("Note totale :")[0].strip()
                # evaluations[f"{complexity}_realism_score"] = realism_eval.split("Note totale :")[-1].strip()

                # # Extract evaluations and scores for standalone quality
                # evaluations[f"{complexity}_standalone"] = standalone_eval.split("Évaluation : ")[-1].split("Note totale :")[0].strip()
                # evaluations[f"{complexity}_standalone_score"] = standalone_eval.split("Note totale :")[-1].strip()

            # Update the output with evaluations
            output.update(evaluations)

        except Exception as e:
            # Handle any errors while reading the document (e.g., file access issues)
            logger.error(f"Error evaluating QA pair: {e}")
            continue

    return outputs  # Return the updated outputs with evaluations

@app.command()
def main(
    eval_corpus_path: Path = Path,  # Default input path
    n_generations: int = 100,  # Number of generations by default, It caps when it reaches the maximum chunks possible from the available documents.
    output_path: Path = Path(__file__).parent / "qa_output.xlsx",  # Default output path
    report_output_path: Path = Path(__file__).parent / "report_output.xlsx"  # Default report output path
):
    """
    Function to generate the synthethic data part of the RAG pipeline.
    Currently supports multiple complexity.
    """
    qa_pairs = generate_qa_pairs(eval_corpus_path, n_generations, output_path)

    if qa_pairs:
        evaluated_pairs = evaluate_qa_pairs(qa_pairs)

        output_data = []
        for output in evaluated_pairs:
            for complexity in ["simple", "reasoning", "multi_context"]:
                output_data.append({
                    "context": output["context"],
                    "question": output["questions"][complexity],
                    "answer": output["answers"][complexity],
                    "complexity": complexity,
                    "source_doc": ", ".join(output["source_docs"]),  # Join multiple sources as a single string
                    "groundedness_evaluation": output.get(f"{complexity}_groundedness"),
                    "groundedness_score": output.get(f"{complexity}_groundedness_score"),
                    # "realism_evaluation": output.get(f"{complexity}_realism"),
                    # "realism_score": output.get(f"{complexity}_realism_score"),
                    # "standalone_evaluation": output.get(f"{complexity}_standalone"),
                    # "standalone_score": output.get(f"{complexity}_standalone_score"),
                })

        # Save the final evaluated QA pairs to an Excel file
        df_output = pd.DataFrame(output_data)
        
        # Apply the function to the score columns
        # df_output['groundedness_score'] = df_output['groundedness_score'].apply(extract_numeric)
        # df_output['realism_score'] = df_output['realism_score'].apply(extract_numeric)
        # df_output['standalone_score'] = df_output['standalone_score'].apply(extract_numeric)
        
        df_output.to_excel(report_output_path, index=False)
        logger.info(f"Final evaluated QA dataset saved to {report_output_path}")

        # Filter the rows based on numeric conditions
        filtered_output = df_output[(df_output['groundedness_score'] >= 4)]

        # filtered_output = df_output[
        #     (df_output['groundedness_score'] >= 4) &
        #     (df_output['realism_score'] >= 4) &
        #     (df_output['standalone_score'] >= 4)
        # ]

        # Save the filtered data to an Excel file
        filtered_output = filtered_output[["context", "question", "answer", "complexity", "source_doc"]]
        filtered_output.to_excel(output_path, index=False)
        logger.info(f"Filtered QA dataset saved to {output_path}")

# Run the app
if __name__ == "__main__":
    app()